"""Generate sample files for Postman / curl testing.

Outputs (overwrites existing files in samples/):
    samples/compliance_handbook.txt
    samples/engineering_design.pdf
    samples/hr_onboarding.docx
    samples/customer_contract.pdf
    samples/unsupported.zip

Re-run with `python scripts/generate_samples.py` from the repo root.
"""
from __future__ import annotations

import io
import zipfile
from pathlib import Path

from docx import Document as DocxDocument
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

OUT = Path(__file__).resolve().parent.parent / "samples"
OUT.mkdir(exist_ok=True)


# -------- compliance handbook (txt) -----------------------------------------
COMPLIANCE_TXT = """Acme Corp Compliance Handbook (Excerpt)

Section 1. Data Retention Policy
All customer records must be retained for a minimum of seven (7) years from the
date of the last transaction. After that period, records may be archived to cold
storage for an additional three (3) years before being permanently destroyed.
Deletion requests under GDPR or CCPA take precedence over the retention schedule
and must be honored within 30 days of receipt.

Section 2. Incident Response
A Severity-1 incident is defined as any event causing total service outage, data
loss affecting more than 1,000 customers, or unauthorized access to production
systems. Severity-1 incidents must be acknowledged within 15 minutes, with a
public status update posted within 60 minutes.

Section 3. Vendor Risk
All third-party vendors handling personal data must complete a SOC 2 Type II
audit annually. Vendors that store data outside of the United States or European
Union require explicit Data Protection Officer approval before onboarding.

Section 4. Approved AI Providers
The following Large Language Model providers are approved for production use:
Anthropic (Claude), OpenAI (GPT-4 family, via Azure only), and Google (Gemini,
via Vertex AI only). Use of any other provider requires written approval from
the Chief Information Security Officer.
"""


def write_txt() -> Path:
    p = OUT / "compliance_handbook.txt"
    p.write_text(COMPLIANCE_TXT)
    return p


# -------- engineering design (pdf) ------------------------------------------
ENGINEERING_PDF_PARAS = [
    ("Title", "Project Atlas — Architecture Design Doc (v1.2)"),
    ("BodyText",
     "Project Atlas is the next-generation document ingestion service for the "
     "Acme platform. It replaces the legacy Mercury pipeline, which suffered "
     "from a 12% job-failure rate and a P95 ingestion latency of 47 seconds "
     "per document. Atlas targets a 0.1% failure rate and P95 of 6 seconds."),
    ("Heading2", "1. Goals"),
    ("BodyText",
     "Atlas must support 50,000 documents per hour at peak, with bursts up to "
     "120,000 per hour for a maximum of 15 minutes. Documents may be PDFs (most "
     "common), DOCX, HTML, plain text, and OCR-scanned PDFs. The system must "
     "tolerate failures of any single embedding provider and any single "
     "vector database shard."),
    ("Heading2", "2. Architecture"),
    ("BodyText",
     "The pipeline is composed of an API layer (FastAPI), an ingestion queue "
     "(NATS JetStream), a worker fleet running on Kubernetes, an embedding "
     "service backed by GPU pods serving Voyage and BGE models, and a Qdrant "
     "cluster sharded by tenant_id hash."),
    ("Heading2", "3. Failure Modes"),
    ("BodyText",
     "If the primary embedding provider returns latency above 800ms or a 5xx "
     "rate above 1% over a 60-second window, the circuit breaker opens and the "
     "fallback provider takes over. If the Qdrant cluster reports any shard as "
     "unhealthy, ingestion is paused for that tenant and queued upstream until "
     "the shard recovers — we never index against a degraded shard."),
    ("Heading2", "4. SLOs"),
    ("BodyText",
     "Atlas commits to: 99.9% availability for the API, 99.5% successful "
     "indexing within 30 seconds of upload, and zero cross-tenant data leaks "
     "(measured by a continuous fuzzer that attempts unauthorized retrievals "
     "against random tenant collections)."),
]


def write_engineering_pdf() -> Path:
    p = OUT / "engineering_design.pdf"
    doc = SimpleDocTemplate(
        str(p), pagesize=LETTER,
        leftMargin=0.7 * inch, rightMargin=0.7 * inch,
        topMargin=0.7 * inch, bottomMargin=0.7 * inch,
    )
    styles = getSampleStyleSheet()
    flow = []
    for style, text in ENGINEERING_PDF_PARAS:
        flow.append(Paragraph(text, styles[style]))
        flow.append(Spacer(1, 0.12 * inch))
    doc.build(flow)
    return p


# -------- HR onboarding (docx) ----------------------------------------------
HR_DOCX_SECTIONS = [
    ("title", "Acme Corp — New Hire Onboarding Guide"),
    ("h1", "Day 1: Equipment & Access"),
    ("p", "All new employees receive a company-issued MacBook Pro (16GB RAM "
          "minimum, 32GB for Engineering) on their start date. Access to "
          "internal systems is provisioned through Okta within 4 business "
          "hours of HR completing the I-9 verification."),
    ("p", "Hardware that fails within the first 30 days is replaced same-day; "
          "after 30 days, the standard 5-business-day replacement SLA applies."),
    ("h1", "Week 1: Required Training"),
    ("p", "Every new hire must complete: Information Security Awareness "
          "(60 min), Anti-Harassment (45 min), and the relevant role-specific "
          "track (Engineering = Code of Conduct + Open Source Policy; Sales = "
          "Anti-Bribery + Data Privacy for Sales; etc.). Training must be "
          "completed within 14 days of start date or system access is revoked."),
    ("h1", "PTO & Holidays"),
    ("p", "Acme observes a flexible PTO policy — there is no fixed accrual; "
          "managers approve based on team coverage. The minimum expected "
          "vacation is 15 business days per calendar year. Acme also closes "
          "for one full week between December 24 and January 1 — this week "
          "does not count against personal PTO."),
    ("h1", "Probation"),
    ("p", "All new hires are subject to a 90-day probationary period. During "
          "this window, performance is reviewed at 30, 60, and 90 days. "
          "Termination during probation does not require the standard "
          "performance improvement plan."),
]


def write_hr_docx() -> Path:
    p = OUT / "hr_onboarding.docx"
    d = DocxDocument()
    for kind, text in HR_DOCX_SECTIONS:
        if kind == "title":
            d.add_heading(text, level=0)
        elif kind == "h1":
            d.add_heading(text, level=1)
        else:
            d.add_paragraph(text)
    d.save(str(p))
    return p


# -------- customer contract (pdf) -------------------------------------------
CONTRACT_PDF_PARAS = [
    ("Title", "Master Services Agreement — Acme & Globex"),
    ("BodyText",
     "This Master Services Agreement ('Agreement') is entered into as of "
     "January 1, 2026, between Acme Corporation ('Acme') and Globex "
     "Industries ('Customer'). The initial term is twenty-four (24) months."),
    ("Heading2", "Section 1. Service Levels"),
    ("BodyText",
     "Acme shall maintain 99.9% monthly uptime for the Production Service. "
     "Downtime exceeding 0.1% in any calendar month entitles Customer to "
     "service credits equal to 10% of monthly fees per 0.1% of additional "
     "downtime, capped at 50% of monthly fees."),
    ("Heading2", "Section 2. Data Processing"),
    ("BodyText",
     "Acme processes Customer Data solely to provide the Service. Customer "
     "Data is encrypted at rest using AES-256 and in transit using TLS 1.3 "
     "or higher. All sub-processors require Customer's prior written approval, "
     "with a list maintained at acme.com/subprocessors."),
    ("Heading2", "Section 3. Term & Termination"),
    ("BodyText",
     "Either party may terminate for material breach with 30 days' written "
     "notice and opportunity to cure. Customer may terminate for convenience "
     "with 90 days' notice; in such case, prepaid fees for services not yet "
     "delivered are refunded pro-rata."),
    ("Heading2", "Section 4. Limitation of Liability"),
    ("BodyText",
     "Each party's aggregate liability under this Agreement is capped at the "
     "fees paid in the twelve (12) months preceding the claim. Neither party "
     "is liable for indirect, consequential, or punitive damages, except for "
     "breaches of confidentiality or data-protection obligations."),
]


def write_contract_pdf() -> Path:
    p = OUT / "customer_contract.pdf"
    doc = SimpleDocTemplate(
        str(p), pagesize=LETTER,
        leftMargin=0.7 * inch, rightMargin=0.7 * inch,
        topMargin=0.7 * inch, bottomMargin=0.7 * inch,
    )
    styles = getSampleStyleSheet()
    flow = []
    for style, text in CONTRACT_PDF_PARAS:
        flow.append(Paragraph(text, styles[style]))
        flow.append(Spacer(1, 0.12 * inch))
    doc.build(flow)
    return p


# -------- unsupported file (zip) --------------------------------------------
def write_unsupported_zip() -> Path:
    p = OUT / "unsupported.zip"
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("readme.txt", "This file exists to test the 415 path.")
    p.write_bytes(buf.getvalue())
    return p


def main() -> None:
    paths = [
        write_txt(),
        write_engineering_pdf(),
        write_hr_docx(),
        write_contract_pdf(),
        write_unsupported_zip(),
    ]
    for p in paths:
        print(f"  {p.relative_to(p.parent.parent)}  ({p.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
