"""Text extraction for supported formats.

MVP supports PDF, DOCX, plain text. HTML and OCR-enabled scanned PDFs are
documented in TRADEOFFS.md as the next addition (BeautifulSoup + Tesseract).
"""
import io

from docx import Document as DocxDocument
from pypdf import PdfReader

SUPPORTED_MIME_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "text",
    "text/markdown": "text",
}


def extract_text(payload: bytes, mime_type: str) -> str:
    kind = SUPPORTED_MIME_TYPES.get(mime_type)
    if kind == "pdf":
        return _extract_pdf(payload)
    if kind == "docx":
        return _extract_docx(payload)
    if kind == "text":
        return payload.decode("utf-8", errors="replace")
    raise ValueError(f"Unsupported mime type: {mime_type}")


def _extract_pdf(payload: bytes) -> str:
    reader = PdfReader(io.BytesIO(payload))
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            # Skip unreadable pages — typical for scanned PDFs without OCR.
            continue
    return "\n\n".join(p for p in parts if p.strip())


def _extract_docx(payload: bytes) -> str:
    doc = DocxDocument(io.BytesIO(payload))
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells if cell.text)
            if row_text:
                parts.append(row_text)
    return "\n\n".join(parts)
